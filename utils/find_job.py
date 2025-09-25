# -*- coding: utf-8 -*-

"""
@author: lucas
@Function:
@file: find_job.py
@time: 2021/11/3 3:39 下午
"""
import os

import docx
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import LAParams

from pdfminer.converter import PDFPageAggregator


def get_files(path):
    res = []
    for i in os.listdir(path):
        # 去掉临时文件
        if os.path.isfile(path + i) and '~$' not in i and '.DS' not in i:
            # 去重 1.doc 和 1.docx
            if (path + i).split(".")[0] not in str(res):
                res.append(path + i)
    return res


def pdf_reader(file):
    fp = open(file, "rb")
    # 创建一个与文档相关联的解释器
    parser = PDFParser(fp)
    # PDF文档对象
    doc = PDFDocument(parser)
    # 链接解释器和文档对象
    parser.set_document(doc)
    # doc.set_paeser(parser)
    # 初始化文档
    # doc.initialize("")
    # 创建PDF资源管理器
    resource = PDFResourceManager()
    # 参数分析器
    laparam = LAParams()
    # 创建一个聚合器
    device = PDFPageAggregator(resource, laparams=laparam)
    # 创建PDF页面解释器
    interpreter = PDFPageInterpreter(resource, device)
    # 使用文档对象得到页面集合
    res = ''
    for page in PDFPage.create_pages(doc):
        # 使用页面解释器来读取
        interpreter.process_page(page)
        # 使用聚合器来获取内容
        layout = device.get_result()
        for out in layout:
            if hasattr(out, "get_text"):
                res = res + '' + out.get_text()
    return res


def word_reader(file):
    try:
        # docx 直接读
        if 'docx' in file:
            res = ''
            f = docx.Document(file)
            for para in f.paragraphs:
                res = res + '\n' + para.text
        else:
            # 先转格式doc>docx
            os.system("textutil -convert docx '%s'" % file)
            word_reader(file + 'x')
            res = ''
            f = docx.Document(file + 'x')
            for para in f.paragraphs:
                res = res + '\n' + para.text
        return res
    except:
        # print(file, 'read failed')
        return ''


def file_reader(file):
    if 'doc' in file:
        res = word_reader(file)
    elif 'pdf' in file:
        res = pdf_reader(file)
    else:
        res = '不是doc，也不是pdf，文件格式不支持！'
    return res


if __name__ == '__main__':
    path = "/Users/leiyuxing/job/"
    abs_files = get_files(path)
    print(abs_files)
    for file in abs_files:
        file_text = file_reader(file)
        print(file_text)
